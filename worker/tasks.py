from worker.broker import broker
from worker.media.transcriber import TranscriberFactory
from worker.ocr import OCRConfigError, OCRDisabledError, run_image_ocr, run_pdf_ocr
from shared.logger import worker_log
from shared.database import AsyncSessionLocal
from shared.models import ItemORM
from hub.core.media.downloader import YTDlpService
from hub.core.llm.summarizer import summarize_video_transcript
from hub.core.notifications import create_notification_for_item_users
from hub.core.model_settings import get_model_assignments
from hub.core.llm.service import LLMManager
from hub.core.billing import QuotaExceededException
from shared.time_utils import now_in_app_timezone_naive
from sqlalchemy import select
import os
import json
import asyncio
import shutil
import tempfile
import random
import mimetypes
import httpx
import posixpath
import re
from pathlib import Path
from urllib.parse import quote, unquote, urlparse
from shared.config import ConfigManager, ConfigKeys
from shared.credentials import get_user_credential, resolve_cookie_file_path
from shared.locale import normalize_preferred_locale


def _resolve_bilibili_cookiefile(url: str, user_id: str | None) -> str | None:
    if not url or "bilibili" not in url:
        return None
    if not user_id:
        return None
    return resolve_cookie_file_path(user_id, "bilibili")


async def _mark_item_failed_with_circuit_breaker(
    item_id: str,
    *,
    error_message: str,
    title: str | None = None,
    user_id: str | None = None,
    extra_metadata: dict | None = None,
) -> None:
    worker_log.error(f"❌ [CIRCUIT BREAKER] {item_id} | {error_message}")

    async with AsyncSessionLocal() as session:
        result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
        item = result.scalar_one_or_none()
        if item:
            metadata = dict(item.metadata_extra or {})
            metadata["processing_status"] = "failed"
            metadata["processing_error"] = error_message
            if extra_metadata:
                metadata.update(extra_metadata)
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(
                    metadata_extra=metadata,
                    updated_at=now_in_app_timezone_naive(),
                )
            )
            await session.commit()
            if not title:
                title = item.title

    await create_notification_for_item_users(
        item_id,
        type="error",
        category="upload_processing",
        title="API circuit breaker triggered",
        description=f"Background processing stopped for {title or 'this item'} because the quota was exhausted.",
        preferred_user_id=user_id,
    )


def _parse_github_repo_ref(item: ItemORM) -> tuple[str, str] | None:
    metadata = item.metadata_extra or {}
    owner = metadata.get("owner")
    repo = metadata.get("repo")
    if owner and repo:
        return str(owner), str(repo)

    source_url = item.raw_link or ""
    if "github.com/" not in source_url:
        return None

    try:
        tail = source_url.split("github.com/", 1)[1]
        parts = [part for part in tail.split("/") if part]
        if len(parts) < 2:
            return None
        return parts[0], parts[1].removesuffix(".git")
    except Exception:
        return None


async def _get_github_access_token(user_id: str | None = None) -> str | None:
    if user_id:
        credential = await get_user_credential(user_id, "github")
        if credential:
            return credential.token_value
    token = await ConfigManager.get_config("github_stars", ConfigKeys.TOKEN, user_id=user_id)
    if token:
        return token
    if user_id:
        return await ConfigManager.get_config("github_stars", ConfigKeys.TOKEN)
    return None


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("DOCX parsing dependency python-docx is not installed") from exc

    document = Document(str(file_path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell_text
                for cell in row.cells
                if (cell_text := (cell.text or "").strip())
            ).strip()
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


async def _download_github_readme(
    item: ItemORM,
    output_dir: Path,
    user_id: str | None = None,
) -> tuple[str, str, str]:
    repo_ref = _parse_github_repo_ref(item)
    if not repo_ref:
        raise ValueError("Unable to resolve GitHub repository information from the item")

    owner, repo = repo_ref
    github_token = await _get_github_access_token(user_id)
    headers = {
        "Accept": "application/vnd.github+json",
        "User-Agent": "Pekno-Hub/1.0",
    }
    if github_token:
        headers["Authorization"] = f"Bearer {github_token}"
    readme_url = f"https://api.github.com/repos/{owner}/{repo}/readme"

    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
        response = await client.get(readme_url, headers=headers)
        response.raise_for_status()
        payload = response.json()
        readme_path = str(payload.get("path") or "README.md")
        download_url = payload.get("download_url")
        if download_url:
            raw_headers = {"User-Agent": "Pekno-Hub/1.0"}
            if github_token:
                raw_headers["Authorization"] = f"Bearer {github_token}"
            raw_response = await client.get(download_url, headers=raw_headers)
            raw_response.raise_for_status()
            readme_text = raw_response.text
        else:
            import base64

            encoded_content = str(payload.get("content") or "")
            readme_text = base64.b64decode(encoded_content).decode("utf-8")

    output_path = output_dir / "README.md"
    output_path.write_text(readme_text, encoding="utf-8")
    return str(output_path), readme_text, readme_path


def _is_external_asset_url(url: str) -> bool:
    normalized = (url or "").strip().lower()
    return (
        normalized.startswith("http://")
        or normalized.startswith("https://")
        or normalized.startswith("//")
        or normalized.startswith("data:")
    )


def _extract_asset_path_token(asset_url: str) -> str | None:
    cleaned = (asset_url or "").strip().strip("<>")
    if not cleaned:
        return None

    cleaned = cleaned.replace("\\", "/")
    title_match = re.match(r'^(.*?)(?:\s+["\'][^"\']*["\'])?$', cleaned)
    token = (title_match.group(1) if title_match else cleaned).strip()
    return token or None


def _normalize_repo_relative_paths(readme_repo_path: str, asset_url: str) -> list[str]:
    cleaned = _extract_asset_path_token(asset_url)
    if not cleaned or _is_external_asset_url(cleaned):
        return []

    parsed = urlparse(cleaned)
    asset_path = unquote(parsed.path or "").replace("\\", "/")
    if not asset_path:
        return []

    readme_dir = posixpath.dirname(readme_repo_path or "README.md")
    candidates: list[str] = []

    def _push(candidate: str):
        normalized = posixpath.normpath(candidate).lstrip("/")
        if normalized and not normalized.startswith("..") and normalized not in candidates:
            candidates.append(normalized)

    if asset_path.startswith("/"):
        _push(asset_path.lstrip("/"))
        return candidates

    _push(posixpath.join(readme_dir, asset_path))
    if not asset_path.startswith("../"):
        _push(asset_path.lstrip("./"))

    return candidates


def _build_github_raw_asset_url(owner: str, repo: str, repo_relative_path: str) -> str:
    safe_path = quote(repo_relative_path, safe="/")
    return f"https://raw.githubusercontent.com/{owner}/{repo}/HEAD/{safe_path}"


async def _download_github_asset(
    client: httpx.AsyncClient,
    owner: str,
    repo: str,
    repo_relative_path: str,
    target_path: Path,
) -> bool:
    asset_url = _build_github_raw_asset_url(owner, repo, repo_relative_path)
    try:
        response = await client.get(asset_url, headers={"User-Agent": "Pekno-Hub/1.0"})
        response.raise_for_status()
    except Exception:
        return False

    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(response.content)
    return True


async def _localize_github_readme_assets(
    item: ItemORM,
    output_dir: Path,
    readme_text: str,
    readme_repo_path: str,
) -> tuple[str, list[str]]:
    repo_ref = _parse_github_repo_ref(item)
    if not repo_ref:
        return readme_text, []

    owner, repo = repo_ref
    assets_dir = output_dir / "assets"
    rewritten_markdown = readme_text
    replacement_map: dict[str, str] = {}
    failed_assets: list[str] = []

    markdown_matches = re.findall(r'!\[[^\]]*\]\(([^)]+)\)', readme_text)
    html_matches = re.findall(r'<img\b[^>]*src=["\']([^"\']+)["\']', readme_text, flags=re.IGNORECASE)
    html_srcset_matches = re.findall(
        r'\bsrcset=["\']([^"\']+)["\']',
        readme_text,
        flags=re.IGNORECASE,
    )
    srcset_candidates: list[str] = []
    for srcset_value in html_srcset_matches:
        for candidate in srcset_value.split(","):
            candidate = candidate.strip()
            if not candidate:
                continue
            url_part = candidate.split()[0].strip()
            if url_part:
                srcset_candidates.append(url_part)

    asset_candidates = list(
        dict.fromkeys([*(markdown_matches or []), *(html_matches or []), *srcset_candidates])
    )

    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
        for asset_url in asset_candidates:
            if _is_external_asset_url(asset_url):
                continue

            repo_relative_paths = _normalize_repo_relative_paths(readme_repo_path, asset_url)
            if not repo_relative_paths:
                continue

            downloaded = False
            for repo_relative_path in repo_relative_paths:
                local_target_path = assets_dir / Path(repo_relative_path)
                encoded_repo_relative_path = quote(repo_relative_path, safe="/")
                public_url = f"/api/static/vault/{item.id}/assets/{encoded_repo_relative_path}"
                if await _download_github_asset(client, owner, repo, repo_relative_path, local_target_path):
                    replacement_map[asset_url] = public_url
                    downloaded = True
                    break

            if not downloaded:
                failed_assets.append(asset_url)
                worker_log.warning(
                    f"⚠️ README asset download failed: {asset_url} | tried={repo_relative_paths}"
                )

    for original, rewritten in replacement_map.items():
        rewritten_markdown = rewritten_markdown.replace(f"({original})", f"({rewritten})")
        rewritten_markdown = re.sub(
            rf'(<img\b[^>]*src=["\']){re.escape(original)}(["\'])',
            rf"\1{rewritten}\2",
            rewritten_markdown,
            flags=re.IGNORECASE,
        )

    def _rewrite_srcset_value(match: re.Match[str]) -> str:
        original_value = match.group(1)
        rewritten_candidates: list[str] = []
        for candidate in original_value.split(","):
            raw_candidate = candidate.strip()
            if not raw_candidate:
                continue

            parts = raw_candidate.split()
            asset_url = parts[0].strip()
            descriptor = " ".join(parts[1:]).strip()
            rewritten_url = replacement_map.get(asset_url, asset_url)
            rewritten_candidates.append(
                f"{rewritten_url} {descriptor}".strip()
            )

        rewritten_value = ", ".join(rewritten_candidates) if rewritten_candidates else original_value
        return f'srcset="{rewritten_value}"'

    rewritten_markdown = re.sub(
        r'\bsrcset=["\']([^"\']+)["\']',
        _rewrite_srcset_value,
        rewritten_markdown,
        flags=re.IGNORECASE,
    )

    return rewritten_markdown, failed_assets


def _format_seconds_for_summary(seconds: float) -> str:
    total_seconds = max(0, int(seconds))
    minutes, secs = divmod(total_seconds, 60)
    hours, minutes = divmod(minutes, 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _build_timed_transcript(segments: list[dict]) -> str:
    lines: list[str] = []
    for segment in segments:
        text = str(segment.get("text", "")).strip()
        if not text:
            continue
        start = float(segment.get("start", 0))
        lines.append(f"[{_format_seconds_for_summary(start)}] {text}")
    return "\n".join(lines)


def _build_plain_transcript(segments: list[dict]) -> str:
    return "\n".join(
        str(segment.get("text", "")).strip()
        for segment in (segments or [])
        if str(segment.get("text", "")).strip()
    )


def _resolve_item_locale(metadata: dict | None, preferred_locale: str | None = None) -> str:
    return normalize_preferred_locale(preferred_locale or (metadata or {}).get("preferred_locale"))


def _fallback_long_summary(item_ctx, reason: str, preferred_locale: str | None = None) -> str:
    base_text = (item_ctx.summary or "").strip()
    if not base_text:
        base_text = (
            "The original summary is empty, so there is no fallback reference text."
            if _resolve_item_locale(item_ctx.metadata_extra, preferred_locale) == "en"
            else "原始摘要为空，因此没有可回退的参考文本。"
        )
    if _resolve_item_locale(item_ctx.metadata_extra, preferred_locale) == "zh-CN":
        return (
            "后处理阶段失败。输入可能是纯音乐、环境音，或属于较难稳定识别的内容，因此未生成新的长篇 AI 总结。\n\n"
            "以下保留原始简介或短摘要作为参考：\n\n"
            f"{base_text}"
        )
    return (
        f"Post-processing failed. The input may be pure audio, background music, or otherwise difficult "
        f"to recognize reliably, so no new long-form AI summary was produced.\n\n"
        f"As a fallback, the original description or short summary is preserved below for reference:\n\n"
        f"{base_text}"
    )


def _pick_random_keyframe_timestamps(duration_seconds: float | int | None, count: int = 5) -> list[int]:
    if not duration_seconds or duration_seconds <= 0:
        return []

    max_second = max(1, int(duration_seconds))
    start = 3 if max_second > 6 else 0
    end = max(start, max_second - 3)

    if end <= start:
        return [start]

    candidates = list(range(start, end + 1))
    sample_size = min(count, len(candidates))
    if sample_size <= 0:
        return []

    return sorted(random.sample(candidates, sample_size))


def _sanitize_keyframe_timestamps(
    timestamps: list[int] | None,
    duration_seconds: float | int | None,
) -> list[int]:
    if not timestamps:
        return []

    if not duration_seconds or duration_seconds <= 0:
        return sorted(dict.fromkeys(max(0, int(ts)) for ts in timestamps))

    max_allowed = max(0, int(float(duration_seconds)) - 1)
    sanitized: list[int] = []
    for ts in timestamps:
        normalized = min(max(0, int(ts)), max_allowed)
        if normalized not in sanitized:
            sanitized.append(normalized)
    return sorted(sanitized)


@broker.task(task_name="process_multimedia")
async def process_multimedia_task(
    item_id: str,
    url: str,
    user_id: str | None = None,
    preferred_locale: str | None = None,
):
    """
    终端多媒体 Worker 流水线：
    YTDlp 剥离音轨 -> Whisper 识别 JSON 阵列 -> AI 导演对峙转录偏差生成报告 -> 数据库双写合并
    """
    worker_log.info(f"🎬 Starting multimedia processing task: {item_id} | source={url}")
    
    # 获取 ItemORM 记录供参考
    async with AsyncSessionLocal() as session:
        result = await session.execute(ItemORM.__table__.select().where(ItemORM.id == item_id))
        item_data = result.fetchone()
        if not item_data:
            worker_log.error(f"❌ Item not found for multimedia task: {item_id}")
            return
            
        class TranscribeContext:
            _data = item_data
            @property
            def title(self): return self._data.title
            @property
            def summary(self): return self._data.content_text or self._data.summary
            @property
            def metadata_extra(self): return self._data.metadata_extra or {}
            
        item_ctx = TranscribeContext()
    
    cache_root = Path("data") / "cache" / "media"
    cache_root.mkdir(parents=True, exist_ok=True)
    task_cache_dir = Path(tempfile.mkdtemp(prefix=f"{item_id}_", dir=str(cache_root)))
    requested_audio_path = task_cache_dir / "audio.mp3"
    requested_video_path = task_cache_dir / "video.mp4"
    
    try:
        metadata_extra = item_ctx.metadata_extra
        effective_locale = _resolve_item_locale(metadata_extra, preferred_locale)
        local_asset_path = item_data.local_asset_path
        duration_seconds = metadata_extra.get("duration")

        if local_asset_path and Path(local_asset_path).exists():
            local_source = Path(local_asset_path)
            if not duration_seconds:
                duration_seconds = metadata_extra.get("duration")

            if item_data.intent == "audio":
                audio_path = str(local_source)
                worker_log.info(f"🎧 Using Vault local audio source: {audio_path}")
            else:
                video_path = str(local_source)
                worker_log.info(f"📹 Using Vault local video source: {video_path}")
                cmd = [
                    "ffmpeg", "-y",
                    "-i", video_path,
                    "-vn",
                    "-acodec", "libmp3lame",
                    str(requested_audio_path),
                ]
                proc = await asyncio.create_subprocess_exec(
                    *cmd,
                    stdout=asyncio.subprocess.DEVNULL,
                    stderr=asyncio.subprocess.PIPE,
                )
                _, stderr = await proc.communicate()
                if proc.returncode != 0:
                    err_text = stderr.decode("utf-8", errors="ignore").strip()
                    raise RuntimeError(f"FFmpeg 提取本地视频音轨失败: {err_text[:500]}")
                audio_path = str(requested_audio_path)
                worker_log.info(f"🎧 Extracted audio track from local video: {audio_path}")
        else:
            cookiefile = _resolve_bilibili_cookiefile(url, user_id)
            if cookiefile:
                worker_log.info(f"🍪 Using cookie file for {url}: {cookiefile}")
            stream_info = await asyncio.to_thread(YTDlpService.extract_info, url, {}, cookiefile=cookiefile)
            duration_seconds = stream_info.get("duration")

            worker_log.info(f"📥 Downloading and localizing best-quality media stream with YTDlp... cache_dir={task_cache_dir}")
            audio_path = await asyncio.to_thread(YTDlpService.download_audio, url, str(requested_audio_path), cookiefile=cookiefile)
            worker_log.info(f"🎧 Audio track stored locally: {audio_path}")
            
        # 定位底层 Whisper 挂载型号
        assignments = await get_model_assignments()
        speech_cfg = next((a for a in assignments if a["key"] == "speech_to_text"), None)
        provider = speech_cfg["provider"] if speech_cfg else "local_whisper"
        model_name = speech_cfg["model"] if speech_cfg else "small"
        
        worker_log.info(f"🎙️ Inference layer engaged with {provider}/{model_name}...")
        transcript_result = await TranscriberFactory.transcribe_audio(audio_path, provider, model_name)
        segments = transcript_result.get("segments", [])
        low_confidence = bool(transcript_result.get("low_confidence"))
        language_probability = float(transcript_result.get("language_probability", 0.0) or 0.0)
        
        if not segments and not low_confidence:
            worker_log.warning("⚠️ Silent media detected: no meaningful transcript segments were captured.")
            return

        raw_transcript_json = json.dumps(segments, ensure_ascii=False)
        timed_transcript_text = _build_timed_transcript(segments)
        plain_transcript_text = _build_plain_transcript(segments)

        if low_confidence:
            worker_log.warning(
                "⚠️ Audio language confidence is too low. Falling back to the original summary instead of generating a long-form summary."
            )

            class FallbackSummaryResult:
                def __init__(self, summary: str, keyframe_timestamps: list[int]):
                    self.summary = summary
                    self.keyframe_timestamps = keyframe_timestamps

            summary_result = FallbackSummaryResult(
                summary=_fallback_long_summary(
                    item_ctx,
                    reason=f"language_probability={language_probability:.2f}",
                    preferred_locale=effective_locale,
                ),
                keyframe_timestamps=_pick_random_keyframe_timestamps(duration_seconds, count=5),
            )
        else:
            worker_log.info("🧠 Generating long-form summary with transcript correction...")
            summary_result = await summarize_video_transcript(
                item_ctx,
                timed_transcript_text,
                preferred_locale=effective_locale,
            )
        
        raw_keyframe_timestamps = list(summary_result.keyframe_timestamps or [])
        sanitized_keyframe_timestamps = _sanitize_keyframe_timestamps(raw_keyframe_timestamps, duration_seconds)
        if sanitized_keyframe_timestamps != raw_keyframe_timestamps:
            worker_log.info(
                f"🛠️ Adjusted keyframe anchors to fit media duration: {raw_keyframe_timestamps} -> {sanitized_keyframe_timestamps}"
            )
        worker_log.info(f"📝 Extracted keyframe anchors: {sanitized_keyframe_timestamps}")
        
        # 4. FFmpeg 截帧落地
        keyframes_urls = []
        if sanitized_keyframe_timestamps:
            worker_log.info("📸 Starting keyframe extraction with FFmpeg...")
            os.makedirs(os.path.join("data", "static", "keyframes"), exist_ok=True)
            try:
                if local_asset_path and Path(local_asset_path).exists() and item_data.intent == "video":
                    video_path = local_asset_path
                    worker_log.info(f"🎞️ Using Vault local video for keyframe extraction: {video_path}")
                else:
                    worker_log.info("📹 Downloading local video cache for keyframe extraction...")
                    video_path = await asyncio.to_thread(
                        YTDlpService.download_video_1080p,
                        url,
                        str(requested_video_path),
                        None,
                        cookiefile,
                    )
                    worker_log.info(f"🎞️ Video cached locally: {video_path}")

                for idx, ts in enumerate(sanitized_keyframe_timestamps):
                    filename = f"{item_id}_{idx}_{ts}.jpg"
                    filepath = os.path.join("data", "static", "keyframes", filename)

                    cmd = [
                        "ffmpeg", "-y",
                        "-ss", str(ts),
                        "-i", video_path,
                        "-frames:v", "1",
                        "-q:v", "2",
                        filepath
                    ]

                    proc = await asyncio.create_subprocess_exec(
                        *cmd,
                        stdout=asyncio.subprocess.DEVNULL,
                        stderr=asyncio.subprocess.PIPE
                    )
                    _, stderr = await proc.communicate()

                    if proc.returncode == 0 and os.path.exists(filepath):
                        keyframes_urls.append(f"/api/static/keyframes/{filename}")
                    else:
                        err_text = stderr.decode("utf-8", errors="ignore").strip()
                        worker_log.warning(
                            f"⚠️ Keyframe extraction failed [timestamp={ts}] returncode={proc.returncode}: {err_text[:500]}"
                        )

                worker_log.info(f"✅ Successfully generated {len(keyframes_urls)} keyframe images")
            except Exception as e:
                worker_log.warning(f"⚠️ Keyframe stream extraction interrupted: {e}")
        
        worker_log.info("💾 Persisting source data and summary into PostgreSQL...")
        ai = LLMManager()
        feature_text = "\n\n".join(
            part for part in [
                f"标题：{item_data.title}" if item_data.title else "",
                plain_transcript_text,
            ]
            if part
        ).strip()
        vector = await ai.get_vector(feature_text) if feature_text else None

        async with AsyncSessionLocal() as session:
            meta = dict(item_data.metadata_extra) if item_data.metadata_extra else {}
            meta["raw_transcript"] = raw_transcript_json
            meta["keyframes"] = keyframes_urls
            meta["keyframe_timestamps"] = sanitized_keyframe_timestamps
            meta["has_long_summary"] = True
            meta["long_summary"] = summary_result.summary
            meta["processing_status"] = "completed"
            meta["preferred_locale"] = effective_locale
            if low_confidence:
                meta["summary_fallback_reason"] = "transcription_low_confidence"
                meta["transcription_language_probability"] = language_probability
            
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(
                    content_text=plain_transcript_text or item_data.content_text,
                    metadata_extra=meta,
                    embedding=vector if vector is not None else item_data.embedding,
                )
            )
            await session.commit()
            worker_log.info("🎯 Multimedia processing pipeline completed successfully.")
        await create_notification_for_item_users(
            item_id,
            type="success",
            category="upload_processing",
            title="Multimedia processing completed",
            description=f"Transcription and summarization finished for {item_data.title or 'this item'}.",
            preferred_user_id=user_id,
        )
            
    except QuotaExceededException as exc:
        await _mark_item_failed_with_circuit_breaker(
            item_id,
            error_message=exc.detail,
            title=item_ctx.title,
            user_id=user_id,
        )
    except Exception as e:
        worker_log.error(f"❌ Multimedia processing pipeline failed: {e}")
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            failed_item = result.scalar_one_or_none()
            if failed_item:
                meta = dict(failed_item.metadata_extra or {})
                meta["processing_status"] = "failed"
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=meta, updated_at=now_in_app_timezone_naive())
                )
                await session.commit()
        await create_notification_for_item_users(
            item_id,
            type="error",
            category="upload_processing",
            title="Multimedia processing failed",
            description=str(e)[:160],
            preferred_user_id=user_id,
        )
    finally:
        if task_cache_dir.exists():
            shutil.rmtree(task_cache_dir, ignore_errors=True)
            worker_log.info(f"🧹 Cleaned up multimedia cache directory: {task_cache_dir}")


async def _download_vault_asset_impl(item_id: str, user_id: str | None = None):
    vault_root = Path("data") / "vault"
    vault_root.mkdir(parents=True, exist_ok=True)

    async with AsyncSessionLocal() as session:
        result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
        item = result.scalar_one_or_none()

    if not item:
        worker_log.error(f"❌ Vault download failed because the item does not exist: {item_id}")
        return

    existing_local_path = item.local_asset_path
    needs_github_readme_backfill = (
        item.source_type == "github_star"
        and (
            not (item.metadata_extra or {}).get("vault_readme_text")
            or not (item.metadata_extra or {}).get("vault_readme_assets_localized")
        )
    )
    if existing_local_path and Path(existing_local_path).exists() and not needs_github_readme_backfill:
        worker_log.info(f"⏭️ Vault local asset already exists, skipping re-download: {item_id}")
        return

    if item.source_type == "upload":
        worker_log.info(f"⏭️ Local upload does not require another download: {item_id}")
        return

    metadata = dict(item.metadata_extra or {})
    metadata["vault_download_status"] = "downloading"
    async with AsyncSessionLocal() as session:
        async with session.begin():
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
            )

    source_url = item.raw_link
    try:
        item_dir = vault_root / item_id
        item_dir.mkdir(parents=True, exist_ok=True)
        local_path: str | None = None

        if item.intent in {"video", "audio"}:
            suffix = ".mp4" if item.intent == "video" else ".mp3"
            output_path = item_dir / f"source{suffix}"
            cookiefile = _resolve_bilibili_cookiefile(source_url, user_id)
            if item.intent == "video":
                local_path = await asyncio.to_thread(YTDlpService.download_video_1080p, source_url, str(output_path), None, cookiefile)
            else:
                local_path = await asyncio.to_thread(YTDlpService.download_audio, source_url, str(output_path), None, cookiefile)
        else:
            if item.source_type == "github_star":
                local_path, readme_text, readme_repo_path = await _download_github_readme(item, item_dir, user_id)
                localized_readme_text, failed_assets = await _localize_github_readme_assets(item, item_dir, readme_text, readme_repo_path)
                Path(local_path).write_text(localized_readme_text, encoding="utf-8")
                metadata["vault_readme_text"] = localized_readme_text
                metadata["vault_download_content_type"] = "text/markdown"
                metadata["vault_readme_assets_localized"] = len(failed_assets) == 0
                metadata["vault_readme_asset_failures"] = failed_assets
                metadata["vault_asset_root"] = "vault"
            else:
                guessed_type, _ = mimetypes.guess_type(source_url)
                ext = Path(source_url).suffix or (
                    ".pdf" if guessed_type == "application/pdf" else
                    ".md" if guessed_type == "text/markdown" else
                    ".html"
                )
                output_path = item_dir / f"source{ext}"
                async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                    response = await client.get(source_url)
                    response.raise_for_status()
                    output_path.write_bytes(response.content)
                local_path = str(output_path)

        metadata["vault_download_status"] = "completed"
        metadata.pop("vault_download_error", None)
        metadata.pop("vault_download_error_context", None)

        async with AsyncSessionLocal() as session:
            async with session.begin():
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(
                        local_asset_path=local_path,
                        content_text=item.content_text,
                        metadata_extra=metadata,
                        updated_at=now_in_app_timezone_naive(),
                    )
                )

        worker_log.info(f"📦 Vault asset download completed: {item_id} -> {local_path}")

        already_processed = bool(
            (item.metadata_extra or {}).get("raw_transcript")
            or (item.metadata_extra or {}).get("has_long_summary")
            or (item.metadata_extra or {}).get("keyframes")
        )

        preferred_locale = metadata.get("preferred_locale")

        if item.intent in {"video", "audio"} and not already_processed:
            await process_multimedia_task.kiq(item_id, source_url, user_id, preferred_locale)
        elif item.intent in {"video", "audio"}:
            worker_log.info(f"⏭️ 已存在多媒体分析结果，跳过重复处理: {item_id}")

        if _is_pdf_item(item):
            ocr_meta = metadata.get("ocr") if isinstance(metadata.get("ocr"), dict) else {}
            if ocr_meta.get("status") not in {"processing", "completed"}:
                await process_pdf_ocr_task.kiq(item_id, user_id, preferred_locale)
    except Exception as e:
        worker_log.error(f"❌ Vault 下载失败: {item_id} | {e}")
        metadata = dict(item.metadata_extra or {})
        metadata["vault_download_status"] = "failed"
        metadata["vault_download_error"] = str(e)
        metadata["vault_download_error_context"] = {
            "source_url": source_url,
            "is_bilibili": "bilibili.com" in (source_url or "") or "b23.tv" in (source_url or "") or "BV" in (source_url or ""),
            "github_user_id": user_id,
        }
        async with AsyncSessionLocal() as session:
            async with session.begin():
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
                )


@broker.task(task_name="download_vault_asset")
async def task_download_vault_asset(item_id: str, user_id: str | None = None):
    await _download_vault_asset_impl(item_id, user_id)


def _build_image_feature_text(title: str, result: dict) -> str:
    parts = [
        f"标题：{title}" if title else "",
        f"图片描述：{result.get('short_caption', '')}" if result.get("short_caption") else "",
        f"场景：{result.get('scene', '')}" if result.get("scene") else "",
        f"OCR：{result.get('ocr_text', '')}" if result.get("ocr_text") else "",
    ]
    objects = result.get("objects") or []
    if objects:
        parts.append("关键元素：" + "、".join(str(obj).strip() for obj in objects if str(obj).strip()))
    tags = result.get("tags") or []
    if tags:
        parts.append("标签：" + "、".join(str(tag).strip() for tag in tags if str(tag).strip()))
    return "\n".join(part for part in parts if part).strip()


def _is_pdf_item(item: ItemORM) -> bool:
    metadata = item.metadata_extra or {}
    mime_type = str(
        metadata.get("mime_type")
        or metadata.get("vault_download_content_type")
        or ""
    ).lower()
    candidate = str(item.local_asset_path or item.raw_link or "").lower()
    return item.intent == "document" and (mime_type == "application/pdf" or candidate.endswith(".pdf"))


def _resolve_image_source_path(item: ItemORM) -> Path | None:
    if item.local_asset_path and Path(item.local_asset_path).exists():
        return Path(item.local_asset_path)

    raw_link = str(item.raw_link or "").strip()
    if not raw_link:
        return None

    if raw_link.startswith("/uploads/"):
        candidate = Path("data") / raw_link.removeprefix("/uploads/")
        if candidate.exists():
            return candidate

    if raw_link.startswith("/api/static/vault/"):
        relative = raw_link.removeprefix("/api/static/vault/")
        candidate = Path("data") / "vault" / relative
        if candidate.exists():
            return candidate

    return None


@broker.task(task_name="process_image_understanding")
async def process_image_understanding_task(
    item_id: str,
    user_id: str | None = None,
    preferred_locale: str | None = None,
):
    worker_log.info(f"🖼️ Starting image understanding task: {item_id}")
    cache_root = Path("data") / "cache" / "images"
    cache_root.mkdir(parents=True, exist_ok=True)
    task_cache_dir = Path(tempfile.mkdtemp(prefix=f"image_{item_id}_", dir=str(cache_root)))

    try:
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if not item:
                worker_log.warning(f"⚠️ Skipping image understanding task because item does not exist: {item_id}")
                return

        metadata = dict(item.metadata_extra or {})
        effective_locale = _resolve_item_locale(metadata, preferred_locale)
        metadata["processing_status"] = "processing"
        if metadata.get("has_long_summary") and metadata.get("image_understanding"):
            worker_log.info(f"⏭️ Skipping image understanding task because results already exist: {item_id}")
            return

        ocr_meta = metadata.get("ocr") if isinstance(metadata.get("ocr"), dict) else {}

        source_path = _resolve_image_source_path(item)
        if source_path is None:
            source_url = str(item.raw_link or "").strip()
            if not source_url:
                raise ValueError("图片条目缺少可读取的本地路径或原始链接")
            target_path = task_cache_dir / "source-image"
            async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                response = await client.get(source_url)
                response.raise_for_status()
                suffix = mimetypes.guess_extension(response.headers.get("content-type", "").split(";")[0].strip()) or ".img"
                target_path = target_path.with_suffix(suffix)
                target_path.write_bytes(response.content)
            source_path = target_path

        image_bytes = source_path.read_bytes()
        mime_type = mimetypes.guess_type(str(source_path))[0] or "image/png"
        ocr_text = str(ocr_meta.get("full_text") or "").strip()

        if ocr_meta.get("status") != "completed":
            try:
                ocr_result, ocr_provider, ocr_model = await run_image_ocr(source_path)
                ocr_text = str(ocr_result.get("full_text") or "").strip()
                metadata["ocr"] = {
                    "kind": "image",
                    "status": "completed",
                    "full_text": ocr_text,
                    "blocks": ocr_result.get("blocks") or [],
                    "provider": ocr_provider,
                    "model": ocr_model,
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": None,
                }
            except OCRDisabledError as exc:
                metadata["ocr"] = {
                    "kind": "image",
                    "status": "disabled",
                    "full_text": "",
                    "blocks": [],
                    "provider": "paddleocr_v5_cpu",
                    "model": "PP-OCRv5",
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": str(exc),
                }
            except (OCRConfigError, Exception) as exc:
                worker_log.warning(f"⚠️ Image OCR failed, falling back to pure vision understanding: {item_id} | {exc}")
                metadata["ocr"] = {
                    "kind": "image",
                    "status": "failed",
                    "full_text": "",
                    "blocks": [],
                    "provider": "paddleocr_v5_cpu",
                    "model": "PP-OCRv5",
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": str(exc),
                }

        ai = LLMManager()
        worker_log.info(f"🧠 Invoking image understanding model: {item_id}")
        result, provider_id, model_name = await ai.understand_image(
            image_bytes,
            mime_type,
            ocr_text=ocr_text,
            preferred_locale=effective_locale,
        )
        feature_text = _build_image_feature_text(item.title, result)
        embed_model_name = await ai.get_embedding_model_name()
        vector = await ai.get_vector(feature_text or item.title)

        short_caption = str(result.get("short_caption") or item.summary or item.title).strip()
        long_summary = str(result.get("detailed_summary_markdown") or short_caption).strip()
        tags = [str(tag).strip() for tag in (result.get("tags") or []) if str(tag).strip()]

        image_understanding_meta = {
            "caption": short_caption,
            "tags": tags,
            "ocr_text": str(result.get("ocr_text") or "").strip(),
            "objects": [str(obj).strip() for obj in (result.get("objects") or []) if str(obj).strip()],
            "scene": str(result.get("scene") or "").strip(),
            "provider": provider_id,
            "model": model_name,
            "embedding_model": embed_model_name,
        }

        metadata["has_long_summary"] = True
        metadata["long_summary"] = long_summary
        metadata["image_understanding"] = image_understanding_meta
        metadata["processing_status"] = "completed"
        metadata["preferred_locale"] = effective_locale
        metadata.pop("image_understanding_error", None)

        async with AsyncSessionLocal() as session:
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(
                    summary=short_caption,
                    content_text=feature_text,
                    tags=tags,
                    metadata_extra=metadata,
                    embedding=vector,
                    updated_at=now_in_app_timezone_naive(),
                )
            )
            await session.commit()
        worker_log.info(f"✅ Image understanding completed: {item_id}")
        await create_notification_for_item_users(
            item_id,
            type="success",
            category="upload_processing",
            title="Image processing completed",
            description=f"Recognition and summarization finished for {item.title or 'this image'}.",
            preferred_user_id=user_id,
        )

    except QuotaExceededException as exc:
        await _mark_item_failed_with_circuit_breaker(
            item_id,
            error_message=exc.detail,
            title=item.title if 'item' in locals() and item else None,
            user_id=user_id,
            extra_metadata={"image_understanding_error": exc.detail},
        )
    except Exception as e:
        worker_log.error(f"❌ Image understanding task failed: {item_id} | {e}")
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if item:
                metadata = dict(item.metadata_extra or {})
                metadata["image_understanding_error"] = str(e)
                metadata["processing_status"] = "failed"
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(
                        metadata_extra=metadata,
                        updated_at=now_in_app_timezone_naive(),
                    )
                )
                await session.commit()
        await create_notification_for_item_users(
            item_id,
            type="error",
            category="upload_processing",
            title="Image processing failed",
            description=str(e)[:160],
            preferred_user_id=user_id,
        )
    finally:
        if task_cache_dir.exists():
            shutil.rmtree(task_cache_dir, ignore_errors=True)


@broker.task(task_name="process_pdf_ocr")
async def process_pdf_ocr_task(
    item_id: str,
    user_id: str | None = None,
    preferred_locale: str | None = None,
):
    worker_log.info(f"📄 Starting PDF OCR task: {item_id}")

    try:
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if not item:
                worker_log.warning(f"⚠️ Skipping PDF OCR because item does not exist: {item_id}")
                return

        if not _is_pdf_item(item):
            worker_log.info(f"⏭️ Skipping PDF OCR because the item is not a PDF: {item_id}")
            return

        if not item.local_asset_path or not Path(item.local_asset_path).exists():
            worker_log.warning(f"⚠️ Skipping PDF OCR because the local file is missing: {item_id}")
            return

        metadata = dict(item.metadata_extra or {})
        effective_locale = _resolve_item_locale(metadata, preferred_locale)
        ocr_meta = metadata.get("ocr") if isinstance(metadata.get("ocr"), dict) else {}
        if ocr_meta.get("status") == "completed" and ocr_meta.get("kind") == "pdf":
            worker_log.info(f"⏭️ Skipping PDF OCR because cached results already exist: {item_id}")
            return

        metadata["ocr"] = {
            "kind": "pdf",
            "status": "processing",
            "full_text": str(ocr_meta.get("full_text") or ""),
            "pages": ocr_meta.get("pages") or [],
            "provider": "paddleocr_v5_cpu",
            "model": "PP-OCRv5",
            "processed_at": now_in_app_timezone_naive().isoformat(),
            "error": None,
        }
        metadata["processing_status"] = "processing"
        async with AsyncSessionLocal() as session:
            async with session.begin():
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
                )

        result, provider_id, model_name = await run_pdf_ocr(item.local_asset_path)
        full_text = str(result.get("full_text") or "").strip()

        metadata["ocr"] = {
            "kind": "pdf",
            "status": "completed",
            "full_text": full_text,
            "pages": result.get("pages") or [],
            "ocr_page_count": int(result.get("ocr_page_count") or 0),
            "provider": provider_id,
            "model": model_name,
            "processed_at": now_in_app_timezone_naive().isoformat(),
            "error": None,
        }
        metadata["processing_status"] = "completed"
        metadata["preferred_locale"] = effective_locale

        values = {
            "metadata_extra": metadata,
            "updated_at": now_in_app_timezone_naive(),
        }

        if full_text:
            ai = LLMManager()
            feature_text = "\n".join(part for part in [f"标题：{item.title}" if item.title else "", full_text] if part).strip()
            vector = await ai.get_vector(feature_text)
            values["content_text"] = full_text
            values["embedding"] = vector

        async with AsyncSessionLocal() as session:
            async with session.begin():
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(**values)
                )

        worker_log.info(f"✅ PDF OCR completed: {item_id}")
        await create_notification_for_item_users(
            item_id,
            type="success",
            category="upload_processing",
            title="PDF OCR completed",
            description=f"Text recognition finished for {item.title or 'this PDF'}.",
            preferred_user_id=user_id,
        )
    except OCRDisabledError as exc:
        worker_log.warning(f"⚠️ PDF OCR is disabled: {item_id}")
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if item:
                metadata = dict(item.metadata_extra or {})
                metadata["ocr"] = {
                    "kind": "pdf",
                    "status": "disabled",
                    "full_text": "",
                    "pages": [],
                    "provider": "paddleocr_v5_cpu",
                    "model": "PP-OCRv5",
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": str(exc),
                }
                metadata["processing_status"] = "disabled"
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
                )
                await session.commit()
    except QuotaExceededException as exc:
        await _mark_item_failed_with_circuit_breaker(
            item_id,
            error_message=exc.detail,
            title=item.title if 'item' in locals() and item else None,
            user_id=user_id,
            extra_metadata={
                "ocr": {
                    "kind": "pdf",
                    "status": "failed",
                    "full_text": "",
                    "pages": [],
                    "provider": "paddleocr_v5_cpu",
                    "model": "PP-OCRv5",
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": exc.detail,
                }
            },
        )
    except Exception as exc:
        worker_log.error(f"❌ PDF OCR task failed: {item_id} | {exc}")
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if item:
                metadata = dict(item.metadata_extra or {})
                metadata["ocr"] = {
                    "kind": "pdf",
                    "status": "failed",
                    "full_text": "",
                    "pages": [],
                    "provider": "paddleocr_v5_cpu",
                    "model": "PP-OCRv5",
                    "processed_at": now_in_app_timezone_naive().isoformat(),
                    "error": str(exc),
                }
                metadata["processing_status"] = "failed"
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
                )
                await session.commit()
        await create_notification_for_item_users(
            item_id,
            type="error",
            category="upload_processing",
            title="PDF OCR failed",
            description=str(exc)[:160],
            preferred_user_id=user_id,
        )


@broker.task(task_name="process_uploaded_text_document")
async def process_uploaded_text_document_task(
    item_id: str,
    user_id: str | None = None,
    preferred_locale: str | None = None,
):
    worker_log.info(f"📝 Starting uploaded text analysis task: {item_id}")
    mime_type = ""

    try:
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if not item:
                worker_log.warning(f"⚠️ Skipping uploaded text analysis because item does not exist: {item_id}")
                return

        metadata = dict(item.metadata_extra or {})
        effective_locale = _resolve_item_locale(metadata, preferred_locale)
        mime_type = str(metadata.get("mime_type") or "").lower()
        if mime_type not in {
            "text/plain",
            "text/markdown",
            "text/x-markdown",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }:
            worker_log.info(f"⏭️ Skipping uploaded text analysis because the item is not TXT/Markdown/DOCX: {item_id}")
            return
        if item.embedding is not None and item.tags:
            worker_log.info(f"⏭️ Skipping uploaded text analysis because vectors and tags already exist: {item_id}")
            return
        if not item.local_asset_path or not Path(item.local_asset_path).exists():
            worker_log.warning(f"⚠️ Skipping uploaded text analysis because the local file is missing: {item_id}")
            return

        metadata["processing_status"] = "processing"
        async with AsyncSessionLocal() as session:
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
            )
            await session.commit()

        source_path = Path(item.local_asset_path)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text_content = _extract_docx_text(source_path)
        else:
            raw_bytes = source_path.read_bytes()
            text_content = ""
            for encoding in ("utf-8-sig", "utf-8", "utf-16", "gb18030"):
                try:
                    text_content = raw_bytes.decode(encoding).strip()
                    break
                except UnicodeDecodeError:
                    continue
            if not text_content:
                text_content = raw_bytes.decode("utf-8", errors="ignore").strip()

        user_supplied_summary = str(metadata.get("user_supplied_summary") or "").strip()
        content_text = "\n\n".join(part for part in [user_supplied_summary, text_content] if part).strip()

        if not content_text:
            raise ValueError("No usable text was extracted from the document.")

        ai = LLMManager()
        short_summary = await ai.generate_summary(
            content_text or item.title,
            length="short",
            preferred_locale=effective_locale,
        )
        tags = await ai.extract_tags(content_text or item.title, preferred_locale=effective_locale)
        feature_text = "\n".join(part for part in [item.title, content_text, " ".join(tags)] if part).strip()
        vector = await ai.get_vector(feature_text or item.title)

        metadata["processing_status"] = "completed"
        metadata["preferred_locale"] = effective_locale
        metadata.pop("text_processing_error", None)
        metadata.pop("docx_extract_error", None)
        async with AsyncSessionLocal() as session:
            await session.execute(
                ItemORM.__table__.update()
                .where(ItemORM.id == item_id)
                .values(
                    content_text=content_text,
                    summary=short_summary or item.summary,
                    tags=tags,
                    embedding=vector,
                    metadata_extra=metadata,
                    updated_at=now_in_app_timezone_naive(),
                )
            )
            await session.commit()

        worker_log.info(f"✅ Uploaded text analysis completed: {item_id}")
        await create_notification_for_item_users(
            item_id,
            type="success",
            category="upload_processing",
            title="Document processing completed",
            description=f"Summarization and indexing finished for {item.title or 'this document'}.",
            preferred_user_id=user_id,
        )
    except QuotaExceededException as exc:
        await _mark_item_failed_with_circuit_breaker(
            item_id,
            error_message=exc.detail,
            title=item.title if 'item' in locals() and item else None,
            user_id=user_id,
            extra_metadata={
                "text_processing_error": exc.detail,
                **(
                    {"docx_extract_error": exc.detail}
                    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else {}
                ),
            },
        )
    except Exception as exc:
        worker_log.error(f"❌ Uploaded text analysis task failed: {item_id} | {exc}")
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(ItemORM).where(ItemORM.id == item_id))
            item = result.scalar_one_or_none()
            if item:
                metadata = dict(item.metadata_extra or {})
                metadata["processing_status"] = "failed"
                metadata["text_processing_error"] = str(exc)
                if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    metadata["docx_extract_error"] = str(exc)
                await session.execute(
                    ItemORM.__table__.update()
                    .where(ItemORM.id == item_id)
                    .values(metadata_extra=metadata, updated_at=now_in_app_timezone_naive())
                )
                await session.commit()
        await create_notification_for_item_users(
            item_id,
            type="error",
            category="upload_processing",
            title="Document processing failed",
            description=str(exc)[:160],
            preferred_user_id=user_id,
        )
